"""Generuje dokumentację projektu Komunikator SR w formacie .docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Style globalne ────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.name = "Calibri"
    return p

def para(text="", bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
    if align:
        p.alignment = align
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def table_header_row(table, headers):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D6EAF8")
        tcPr.append(shd)

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_header_row(t, headers)
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(cell_text)
            for run in c.paragraphs[0].runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
    return t

def image_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ZRZUT EKRANU: {caption}]")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x85, 0x92, 0x9E)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
# STRONA TYTUŁOWA
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("KOMUNIKATOR SR")
run.font.name = "Calibri"
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Projekt zaliczeniowy — Systemy Rozproszone")
run.font.name = "Calibri"
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Dokumentacja techniczna")
run.font.name = "Calibri"
run.font.size = Pt(13)
run.font.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Autorzy:\nDawid Świgut\nKrzysztof Witek\nJakub Zasadni")
run.font.name = "Calibri"
run.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2025 / 2026")
run.font.name = "Calibri"
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. OPIS PROJEKTU
# ══════════════════════════════════════════════════════════════════════════════

heading("1. Opis projektu", 1)
para(
    "Komunikator SR to aplikacja webowa czasu rzeczywistego umożliwiająca prowadzenie "
    "rozmów prywatnych (1:1) oraz grupowych (pokoje). Projekt realizuje zagadnienia "
    "systemów rozproszonych: komunikację asynchroniczną przez WebSocket, zarządzanie "
    "sesjami wielu klientów oraz persystencję danych w relacyjnej bazie."
)
doc.add_paragraph()
para("Kluczowe funkcjonalności:", bold=True)
bullet("Rejestracja i logowanie użytkowników (JWT + bcrypt)")
bullet("Czat prywatny 1:1 z historią wiadomości i wskaźnikiem pisania")
bullet("Pokoje grupowe — tworzenie, dołączanie, zapraszanie, usuwanie")
bullet("Status online/offline użytkowników w czasie rzeczywistym")
bullet("Powiadomienia o nowych wiadomościach (toast + badge + tytuł zakładki)")
bullet("Panel administratora — zarządzanie użytkownikami i pokojami")
bullet("Deployment w chmurze Azure (VM + Docker Compose)")
doc.add_paragraph()
para("Podział odpowiedzialności:", bold=True)
add_table(
    ["Obszar", "Osoba"],
    [
        ["Backend (Flask, Socket.IO, REST API)", "Dawid Świgut"],
        ["Frontend (React, UI/UX)", "Krzysztof Witek"],
        ["Baza danych, DevOps, Infrastruktura", "Jakub Zasadni"],
    ],
)


# ══════════════════════════════════════════════════════════════════════════════
# 2. SCENARIUSZE BIZNESOWE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("2. Scenariusze biznesowe (przypadki użycia)", 1)

para(
    "System obsługuje pięciu aktorów: Gość (niezalogowany), Użytkownik (zalogowany), "
    "Właściciel Pokoju (właściciel konkretnego pokoju), Administrator oraz systemy "
    "zewnętrzne PostgreSQL i Socket.IO."
)
doc.add_paragraph()

heading("2.1. Autentykacja", 2)
add_table(
    ["ID", "Nazwa", "Aktor", "Opis"],
    [
        ["UC01", "Rejestracja", "Gość", "Podanie username, email i hasła. Walidacja unikalności. Hasło hashowane bcrypt. Pierwszy zarejestrowany użytkownik otrzymuje rolę admina."],
        ["UC02", "Logowanie", "Gość", "Weryfikacja email + hasło. Zwrot tokenu JWT (HS256, czas życia: 1h)."],
        ["UC03", "Wylogowanie", "Użytkownik", "Klient usuwa token z localStorage i rozłącza WebSocket."],
        ["UC04", "Weryfikacja JWT", "System", "Każde żądanie do chronionych endpointów wymaga nagłówka Authorization: Bearer <token>."],
    ],
)
doc.add_paragraph()

heading("2.2. Czat prywatny 1:1", 2)
add_table(
    ["ID", "Nazwa", "Aktor", "Opis"],
    [
        ["UC05", "Wysłanie wiadomości prywatnej", "Użytkownik", "Emit zdarzenia private_message przez WebSocket. Wiadomość zapisywana w DB, dostarczana do odbiorcy."],
        ["UC06", "Odbieranie wiadomości", "Użytkownik", "Real-time dostarczenie przez socket. Historia ładowana przy wyborze kontaktu."],
        ["UC07", "Wskaźnik pisania", "Użytkownik", "Emit zdarzenia typing, odbiorca widzi '... pisze' przez 2,5 s."],
        ["UC08", "Potwierdzenie dostarczenia (ACK)", "System", "Po dostarczeniu wiadomości backend emituje ack do nadawcy."],
        ["UC09", "Historia wiadomości", "Użytkownik", "GET /api/messages/<user_id> — ostatnie 50 wiadomości, posortowane rosnąco."],
    ],
)
doc.add_paragraph()

heading("2.3. Pokoje grupowe — uczestnik", 2)
add_table(
    ["ID", "Nazwa", "Aktor", "Opis"],
    [
        ["UC10", "Przeglądanie pokojów", "Użytkownik", "GET /api/rooms/all — lista wszystkich pokojów z flagą is_member i listą członków."],
        ["UC11", "Dołączenie do pokoju", "Użytkownik", "POST /api/rooms/<id>/join + emit join_room przez WebSocket."],
        ["UC12", "Wysłanie wiadomości do pokoju", "Użytkownik", "Emit room_message. Backend broadcastuje do wszystkich w Socket.IO room."],
        ["UC13", "Odbieranie wiadomości grupowej", "Użytkownik", "Broadcast z include_self=True. Wiadomości zawierają sender_username."],
        ["UC14", "Historia pokoju", "Użytkownik", "GET /api/rooms/<id>/messages — ostatnie 50 wiadomości po dołączeniu."],
        ["UC15", "Opuszczenie pokoju (WS)", "Użytkownik", "Emit leave_room — wyjście z Socket.IO room (nie usuwa z DB)."],
    ],
)
doc.add_paragraph()

heading("2.4. Zarządzanie pokojem — właściciel", 2)
add_table(
    ["ID", "Nazwa", "Aktor", "Opis"],
    [
        ["UC16", "Tworzenie pokoju", "Użytkownik", "POST /api/rooms/ — tworzący automatycznie staje się właścicielem i pierwszym członkiem."],
        ["UC17", "Zaproszenie użytkownika", "Właściciel", "POST /api/rooms/<id>/invite — backend dodaje RoomMember i emituje room_invite przez Socket.IO."],
        ["UC18", "Usunięcie członka", "Właściciel", "DELETE /api/rooms/<id>/members/<uid>"],
        ["UC19", "Usunięcie pokoju", "Właściciel / Admin", "DELETE /api/rooms/<id> — kaskadowe usunięcie członków i wiadomości."],
    ],
)
doc.add_paragraph()

heading("2.5. Obecność i powiadomienia", 2)
add_table(
    ["ID", "Nazwa", "Aktor", "Opis"],
    [
        ["UC20", "Status online/offline", "System", "Broadcast user_status przy connect/disconnect. Zielona/szara kropka przy nazwie."],
        ["UC21", "Powiadomienie o wiadomości", "System", "Toast (4 s) + badge w sidebarze + zmiana tytułu zakładki przy nowej wiadomości."],
        ["UC22", "Automatyczny reconnect", "System", "Socket.IO reconnect: Infinity prób, delay 1–10 s. Backend auto-joinuje pokoje przy reconnect."],
    ],
)
doc.add_paragraph()

heading("2.6. Administracja", 2)
add_table(
    ["ID", "Nazwa", "Aktor", "Opis"],
    [
        ["UC23", "Logowanie admina", "Administrator", "Identyczne z UC02. Konto tworzone z env vars ADMIN_* przy pierwszym starcie backendu."],
        ["UC24", "Zarządzanie użytkownikami", "Administrator", "GET /api/admin/users — lista, DELETE /api/admin/users/<id> — usunięcie kaskadowe."],
        ["UC25", "Nadanie/odebranie roli admina", "Administrator", "POST /api/admin/users/<id>/toggle-admin"],
        ["UC26", "Zarządzanie pokojami", "Administrator", "GET /api/admin/rooms, DELETE /api/admin/rooms/<id> — może usunąć każdy pokój."],
    ],
)

image_placeholder("Diagram przypadków użycia (use-case.puml)")


# ══════════════════════════════════════════════════════════════════════════════
# 3. ARCHITEKTURA SYSTEMU
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("3. Architektura systemu", 1)

heading("3.1. Stack technologiczny", 2)
add_table(
    ["Warstwa", "Technologia", "Wersja", "Rola"],
    [
        ["Frontend", "React + Vite", "18 / 5", "SPA — interfejs użytkownika"],
        ["Frontend prod", "nginx", "1.27-alpine", "Serwowanie build + reverse proxy"],
        ["Backend", "Flask", "3.0.3", "REST API + WebSocket server"],
        ["WebSocket", "Flask-SocketIO + eventlet", "5.3.6 / 0.36", "Async I/O, zarządzanie połączeniami"],
        ["Auth", "Flask-JWT-Extended + bcrypt", "4.6.0 / 1.0.1", "Tokeny JWT, hashowanie haseł"],
        ["ORM", "Flask-SQLAlchemy + Alembic", "3.1.1 / 4.0.7", "Modele danych, migracje"],
        ["Baza danych", "PostgreSQL", "16-alpine", "Persystencja wiadomości i użytkowników"],
        ["Konteneryzacja", "Docker + Docker Compose", "latest", "Izolacja środowisk, deployment"],
        ["Chmura", "Azure VM (Ubuntu 24.04, B1s)", "-", "Hosting produkcyjny"],
    ],
)
doc.add_paragraph()

heading("3.2. Diagram komponentów i wdrożenia", 2)
para(
    "Aplikacja składa się z trzech kontenerów Docker zarządzanych przez Docker Compose:"
)
bullet("frontend — nginx serwuje zbudowaną aplikację React i proxy do backendu")
bullet("backend — Flask z eventletem obsługuje REST API oraz WebSocket")
bullet("db — PostgreSQL z named volume pgdata dla persystencji danych")
doc.add_paragraph()
image_placeholder("Diagram komponentów i wdrożenia (component-diagram.puml)")
doc.add_paragraph()

heading("3.3. Przepływ komunikacji", 2)
add_table(
    ["Typ", "Protokół", "Kierunek", "Opis"],
    [
        ["Autentykacja", "HTTP/REST + JSON", "Klient → Backend", "POST /api/auth/login zwraca JWT"],
        ["Dane startowe", "HTTP/REST + JSON", "Klient → Backend", "GET /api/users/, /api/rooms/ po zalogowaniu"],
        ["WebSocket", "Socket.IO (WS)", "Dwukierunkowy", "Wiadomości RT, status, typing, pokoje"],
        ["Baza danych", "PostgreSQL wire protocol", "Backend → DB", "SQLAlchemy ORM przez DATABASE_URL"],
    ],
)


# ══════════════════════════════════════════════════════════════════════════════
# 4. SPECYFICZNE ROZWIĄZANIA — SYSTEM ROZPROSZONY
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("4. Specyficzne rozwiązania w kontekście systemu rozproszonego", 1)

heading("4.1. WebSocket zamiast pollingu", 2)
para(
    "Tradycyjne podejście do komunikacji real-time opiera się na HTTP pollingu "
    "(klient pyta serwer co N sekund). W systemie rozproszonym generuje to "
    "znaczące obciążenie i duże opóźnienia. Komunikator SR wykorzystuje protokół "
    "WebSocket (przez bibliotekę Socket.IO), który utrzymuje stałe, dwukierunkowe "
    "połączenie TCP. Serwer może aktywnie \"pushować\" zdarzenia do klientów "
    "bez potrzeby odpytywania."
)
para("Korzyści:", bold=True)
bullet("Opóźnienie wiadomości < 100 ms w sieci lokalnej")
bullet("Brak zbędnego ruchu HTTP przy braku aktywności")
bullet("Serwer może inicjować komunikację (broadcast, powiadomienia)")

heading("4.2. SessionStore — obsługa wielu sesji", 2)
para(
    "Kluczowym wyzwaniem w rozproszonym czacie jest śledzenie, przez które "
    "połączenie WebSocket dostarczyć wiadomość do danego użytkownika. "
    "Jeden użytkownik może mieć otwarte wiele kart przeglądarki lub urządzeń "
    "— każde tworzy osobne połączenie z unikalnym socket_id."
)
para("Rozwiązanie — SessionStore (singleton, thread-safe):", bold=True)
code_block("_sessions: dict[user_id → set(socket_id_1, socket_id_2, ...)]")
para(
    "Przy każdym zdarzeniu private_message backend pobiera zbiór socket_id "
    "dla danego user_id i dostarcza wiadomość do każdego z nich. "
    "Synchronizacja przez threading.Lock zapewnia bezpieczeństwo przy "
    "równoczesnych połączeniach eventlet."
)
bullet("add_session(user_id, socket_id) — przy connect")
bullet("remove_session(user_id, socket_id) — przy disconnect")
bullet("get_socket_ids(user_id) → set — przy wysyłaniu wiadomości")
bullet("get_online_user_ids() → set — przy pobieraniu listy użytkowników")

heading("4.3. Socket.IO Rooms — broadcast grupowy", 2)
para(
    "Do obsługi pokojów grupowych wykorzystano wbudowany mechanizm Socket.IO Rooms. "
    "Każdy pokój (room) w DB ma odpowiadający mu \"kanał\" w Socket.IO o nazwie "
    "równej str(room_id). Wiadomość wysłana do pokoju trafia do wszystkich "
    "klientów aktualnie zasubskrybowanych na ten kanał."
)
code_block("emit('room_message', payload, to=str(room_id), include_self=True)")
para(
    "include_self=True powoduje, że nadawca też otrzymuje broadcast — "
    "dzięki temu nie ma potrzeby robić optimistic update po stronie frontendu "
    "dla wiadomości pokojowych. Eliminuje to problem duplikacji wiadomości."
)
para("Auto-join przy reconnect:", bold=True)
para(
    "Przy każdym zdarzeniu connect backend automatycznie dodaje socket "
    "do wszystkich Socket.IO rooms, których użytkownik jest członkiem w DB. "
    "Zapewnia to odporność na rozłączenia — po powrocie sieciowym użytkownik "
    "nadal odbiera wiadomości grupowe bez potrzeby ręcznego ponownego dołączania."
)

heading("4.4. Bezstanowy JWT", 2)
para(
    "Autentykacja opiera się na bezstanowych tokenach JWT (JSON Web Token) "
    "podpisanych kluczem HS256. Serwer nie przechowuje sesji użytkowników "
    "w żadnym centralnym magazynie — weryfikacja tokenu odbywa się przez "
    "sprawdzenie podpisu kryptograficznego. Umożliwia to horyzontalne skalowanie "
    "backendu (wiele instancji nie musi synchronizować stanu sesji)."
)
bullet("Token przekazywany jako: nagłówek Authorization: Bearer <token> (REST) lub query param ?token= (WebSocket)")
bullet("Czas życia: 1 godzina (konfigurowalny przez JWT_SECRET_KEY)")
bullet("Po wygaśnięciu: klient otrzymuje 401 i musi ponownie się zalogować")

heading("4.5. Asynchroniczne I/O — eventlet", 2)
para(
    "Flask domyślnie jest synchronicznym frameworkiem WSGI — obsługuje jedno "
    "żądanie na raz w wątku. WebSocket wymaga utrzymania długotrwałych połączeń, "
    "co blokowałoby wszystkie inne klientów. Flask-SocketIO z trybem eventlet "
    "stosuje korutyny (greenthreads) zamiast wątków systemowych: tysiące "
    "połączeń współdzieli niewielką pulę wątków dzięki kooperatywnemu "
    "przełączaniu kontekstu przy operacjach I/O."
)
code_block("socketio = SocketIO(async_mode='eventlet')")

heading("4.6. Reconnect i odporność na awarie", 2)
para("Po stronie klienta (Socket.IO JavaScript):")
code_block("reconnectionAttempts: Infinity\nreconnectionDelay: 1000\nreconnectionDelayMax: 10000")
para(
    "Klient próbuje ponownie połączyć się w nieskończoność z exponential backoff "
    "do 10 sekund. Bez tego ustawienia (domyślnie: 5 prób) po krótkim zaniku "
    "sieci użytkownik traciłby możliwość odbierania wiadomości bez przeładowania "
    "strony."
)


# ══════════════════════════════════════════════════════════════════════════════
# 5. PROJEKT ROZWIĄZANIA
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("5. Projekt rozwiązania", 1)

heading("5.1. Model danych", 2)
add_table(
    ["Tabela", "Kluczowe kolumny", "Opis"],
    [
        ["users", "id, username, email, password_hash, is_admin", "Użytkownicy systemu. is_admin=True dla administratorów."],
        ["rooms", "id, name, created_by (FK users)", "Pokoje grupowe. created_by = właściciel."],
        ["room_members", "id, room_id (FK), user_id (FK), UNIQUE(room_id, user_id)", "Relacja N:M użytkownik–pokój."],
        ["messages", "id, sender_id, recipient_id (nullable), room_id (nullable), content, delivered", "Wiadomości 1:1 (recipient_id) lub grupowe (room_id). Dokładnie jedno z tych pól jest NOT NULL."],
    ],
)
doc.add_paragraph()
image_placeholder("Diagram klas — modele danych (class-diagram.puml)")
doc.add_paragraph()

heading("5.2. REST API", 2)
add_table(
    ["Endpoint", "Metoda", "Auth", "Opis"],
    [
        ["POST /api/auth/register", "POST", "—", "Rejestracja. Body: {username, email, password}"],
        ["POST /api/auth/login", "POST", "—", "Logowanie. Zwraca {access_token, user}"],
        ["GET /api/users/", "GET", "JWT", "Lista użytkowników z flagą online"],
        ["GET /api/rooms/", "GET", "JWT", "Pokoje użytkownika (jest członkiem)"],
        ["GET /api/rooms/all", "GET", "JWT", "Wszystkie pokoje z is_member i members[]"],
        ["POST /api/rooms/", "POST", "JWT", "Utwórz pokój. Body: {name}"],
        ["POST /api/rooms/<id>/join", "POST", "JWT", "Dołącz do pokoju"],
        ["POST /api/rooms/<id>/invite", "POST", "JWT+Owner", "Zaproś użytkownika. Body: {user_id}"],
        ["DELETE /api/rooms/<id>/members/<uid>", "DELETE", "JWT+Owner", "Usuń członka"],
        ["DELETE /api/rooms/<id>", "DELETE", "JWT+Owner/Admin", "Usuń pokój"],
        ["GET /api/rooms/<id>/messages", "GET", "JWT+Member", "Historia wiadomości pokoju"],
        ["GET /api/messages/<user_id>", "GET", "JWT", "Historia wiadomości prywatnych"],
        ["GET /api/admin/users", "GET", "JWT+Admin", "Lista wszystkich użytkowników"],
        ["DELETE /api/admin/users/<id>", "DELETE", "JWT+Admin", "Usuń użytkownika"],
        ["POST /api/admin/users/<id>/toggle-admin", "POST", "JWT+Admin", "Nadaj/odbierz rolę admina"],
        ["GET /api/admin/rooms", "GET", "JWT+Admin", "Lista wszystkich pokojów"],
        ["DELETE /api/admin/rooms/<id>", "DELETE", "JWT+Admin", "Usuń pokój (admin override)"],
    ],
)
doc.add_paragraph()

heading("5.3. Zdarzenia WebSocket (Socket.IO)", 2)
para("Zdarzenia Klient → Serwer:", bold=True)
add_table(
    ["Zdarzenie", "Payload", "Opis"],
    [
        ["connect", "?token=<JWT>", "Nawiązanie połączenia. Backend weryfikuje token, rejestruje sesję, auto-joinuje pokoje."],
        ["disconnect", "—", "Usunięcie sesji. Broadcast user_status offline jeśli brak innych sesji."],
        ["private_message", "{to_user_id, content}", "Wiadomość prywatna. Zapisuje do DB, dostarcza do odbiorcy."],
        ["room_message", "{room_id, content}", "Wiadomość do pokoju. Broadcast do Socket.IO room."],
        ["join_room", "{room_id}", "Subskrypcja na Socket.IO room."],
        ["leave_room", "{room_id}", "Anulowanie subskrypcji."],
        ["typing", "{to_user_id}", "Wskaźnik pisania do użytkownika."],
    ],
)
doc.add_paragraph()
para("Zdarzenia Serwer → Klient:", bold=True)
add_table(
    ["Zdarzenie", "Payload", "Opis"],
    [
        ["message", "{id, sender_id, sender_username, content, ...}", "Nowa wiadomość prywatna."],
        ["room_message", "{id, room_id, sender_id, sender_username, content, ...}", "Nowa wiadomość grupowa."],
        ["user_status", "{user_id, online: bool}", "Zmiana statusu użytkownika. Broadcast do wszystkich."],
        ["ack", "{message_id}", "Potwierdzenie dostarczenia wiadomości prywatnej."],
        ["typing", "{from_user_id}", "Wskaźnik pisania od użytkownika."],
        ["room_invite", "{id, name, created_by}", "Zaproszenie do pokoju. Klient dołącza i aktualizuje sidebar."],
        ["error", "{message}", "Błąd operacji WebSocket."],
    ],
)
doc.add_paragraph()

heading("5.4. Sekwencje — diagramy", 2)
image_placeholder("Diagram sekwencji — logowanie i nawiązanie połączenia WS (sequence-login.puml)")
doc.add_paragraph()
image_placeholder("Diagram sekwencji — wysyłanie wiadomości prywatnej i grupowej (sequence-messaging.puml)")


# ══════════════════════════════════════════════════════════════════════════════
# 6. ŚRODOWISKO URUCHOMIENIOWE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("6. Środowisko uruchomieniowe", 1)

heading("6.1. Uruchomienie lokalne (deweloperskie)", 2)
code_block("docker compose up -d --build")
para("Dostęp: http://localhost:5173 (frontend), http://localhost:5000 (backend API)")
para("Domyślne konto admina (z backend/.env):")
bullet("Email: admin@komunikator.local")
bullet("Hasło: admin123")
doc.add_paragraph()

heading("6.2. Deployment produkcyjny (Azure VM)", 2)
para("Środowisko produkcyjne działa na maszynie wirtualnej Azure (Ubuntu 24.04, B1s).")
para("Adres produkcyjny: http://20.56.129.98/ | Panel admina: http://20.56.129.98/admin")
doc.add_paragraph()
para("Architektura produkcyjna:", bold=True)
bullet("Frontend: multi-stage Docker build (npm run build → nginx:1.27-alpine)")
bullet("Backend: Flask + eventlet na porcie 5000 (niedostępny z zewnątrz)")
bullet("nginx: jedyny punkt wejścia (port 80), proxy /api/ i /socket.io/ do backendu")
bullet("Baza danych: PostgreSQL z named volume pgdata (dane przeżywają restart)")
doc.add_paragraph()
code_block(
    "docker compose -f docker-compose.prod.yml --env-file .env.prod up -d --build"
)
doc.add_paragraph()
add_table(
    ["Zmienna (.env.prod)", "Opis"],
    [
        ["DB_PASSWORD", "Hasło do PostgreSQL"],
        ["SECRET_KEY", "Klucz Flask (min. 32 znaki)"],
        ["JWT_SECRET_KEY", "Klucz podpisu tokenów JWT"],
        ["ADMIN_USERNAME", "Login admina (domyślnie: admin)"],
        ["ADMIN_EMAIL", "Email admina"],
        ["ADMIN_PASSWORD", "Hasło admina — tworzone przy starcie jeśli nie istnieje"],
    ],
)


# ══════════════════════════════════════════════════════════════════════════════
# 7. TESTY
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("7. Testy", 1)

para(
    "Testy przeprowadzono manualnie poprzez uruchomienie aplikacji lokalnie "
    "oraz w środowisku produkcyjnym Azure. Poniżej przedstawiono scenariusze "
    "testowe z wynikami."
)
doc.add_paragraph()

heading("7.1. Testy autentykacji", 2)
add_table(
    ["Test", "Kroki", "Oczekiwany wynik", "Wynik"],
    [
        ["Rejestracja nowego użytkownika", "Podaj username, email, hasło → Kliknij Zarejestruj", "Konto utworzone, przekierowanie na login", "✅"],
        ["Rejestracja z istniejącym emailem", "Podaj email już w bazie", "Błąd 409 'Email already in use'", "✅"],
        ["Logowanie poprawnymi danymi", "Podaj email + hasło → Kliknij Zaloguj", "JWT w localStorage, wejście do czatu", "✅"],
        ["Logowanie złym hasłem", "Podaj błędne hasło", "Błąd 401 'Invalid credentials'", "✅"],
        ["Dostęp bez tokenu", "Wywołaj GET /api/users/ bez nagłówka", "Odpowiedź 401 Unauthorized", "✅"],
    ],
)
image_placeholder("Ekran logowania")
image_placeholder("Ekran rejestracji")
doc.add_paragraph()

heading("7.2. Testy czatu prywatnego 1:1", 2)
add_table(
    ["Test", "Kroki", "Oczekiwany wynik", "Wynik"],
    [
        ["Wysłanie wiadomości prywatnej", "Zaloguj dwóch użytkowników → A wysyła wiadomość do B", "B odbiera wiadomość natychmiast", "✅"],
        ["Historia wiadomości", "Wyloguj się i zaloguj ponownie → otwórz czat z kontaktem", "Historia ostatnich 50 wiadomości widoczna", "✅"],
        ["Wskaźnik pisania", "A zaczyna pisać", "B widzi '... pisze' przez 2,5 s", "✅"],
        ["Wielosesyjność", "Zaloguj tego samego użytkownika w 2 kartach", "Wiadomość dociera do obu kart", "✅"],
        ["Wiadomość gdy odbiorca offline", "Wyloguj B → A wysyła wiadomość → zaloguj B", "B widzi wiadomość w historii po zalogowaniu", "✅"],
    ],
)
image_placeholder("Czat prywatny — wiadomości (widok nadawcy i odbiorcy)")
image_placeholder("Wskaźnik pisania")
doc.add_paragraph()

heading("7.3. Testy pokojów grupowych", 2)
add_table(
    ["Test", "Kroki", "Oczekiwany wynik", "Wynik"],
    [
        ["Tworzenie pokoju", "Kliknij '+' → wpisz nazwę → Enter", "Pokój widoczny w sidebarze, automatyczne dołączenie", "✅"],
        ["Wiadomość grupowa", "A i B w tym samym pokoju → A wysyła wiadomość", "B widzi wiadomość z username A", "✅"],
        ["Dołączenie do pokoju", "Kliknij 'Przeglądaj pokoje' → Dołącz", "Pokój dodany do sidebaru, historia widoczna", "✅"],
        ["Zaproszenie przez właściciela", "Owner kliknie 'Zarządzaj' → Zaproś użytkownika", "Użytkownik otrzymuje room_invite, pokój pojawia się w sidebarze", "✅"],
        ["Usunięcie pokoju przez admina", "Admin Panel → Pokoje → Usuń", "Pokój usunięty, członkowie go nie widzą po odświeżeniu", "✅"],
    ],
)
image_placeholder("Pokój grupowy — widok czatu z username nadawcy")
image_placeholder("Modal 'Zarządzaj pokojem'")
image_placeholder("Modal 'Przeglądaj pokoje' z listą członków")
doc.add_paragraph()

heading("7.4. Testy statusu online i powiadomień", 2)
add_table(
    ["Test", "Kroki", "Oczekiwany wynik", "Wynik"],
    [
        ["Status online", "Zaloguj użytkownika B", "Przy A zmienia się kropka B na zieloną natychmiast", "✅"],
        ["Status offline", "Wyloguj B lub zamknij kartę", "Przy A kropka B zmienia się na szarą", "✅"],
        ["Toast powiadomienie", "B wysyła wiadomość gdy A ma otwarty inny czat", "Toast z treścią wiadomości przez 4 s", "✅"],
        ["Badge nieodczytanych", "B wysyła 3 wiadomości gdy A nie ma otwartego czatu z B", "Cyfra '3' przy nazwie B w sidebarze", "✅"],
        ["Tytuł zakładki", "Nieodczytane wiadomości", "Tytuł zmienia się na '(N) Komunikator'", "✅"],
        ["Reconnect", "Tymczasowo wyłącz internet → przywróć", "Połączenie odnawiać się automatycznie, brak potrzeby przelogowania", "✅"],
    ],
)
image_placeholder("Sidebar — zielone/szare kropki statusu online")
image_placeholder("Toast powiadomienie o nowej wiadomości")
doc.add_paragraph()

heading("7.5. Testy panelu admina", 2)
add_table(
    ["Test", "Kroki", "Oczekiwany wynik", "Wynik"],
    [
        ["Dostęp do panelu", "Zaloguj jako admin → kliknij 'Panel admina'", "Strona /admin widoczna z listą użytkowników i pokojów", "✅"],
        ["Brak dostępu dla non-admin", "Zaloguj jako zwykły user → wejdź na /admin", "Przekierowanie na /", "✅"],
        ["Nadanie roli admina", "Panel admina → Użytkownicy → Nadaj admina", "Rola zmieniona, badge 'Admin' przy użytkowniku", "✅"],
        ["Usunięcie użytkownika", "Panel admina → Użytkownicy → Usuń", "Użytkownik usunięty wraz z wiadomościami", "✅"],
        ["Usunięcie pokoju przez admina", "Panel admina → Pokoje → Usuń", "Pokój usunięty", "✅"],
    ],
)
image_placeholder("Panel admina — zakładka Użytkownicy")
image_placeholder("Panel admina — zakładka Pokoje")
doc.add_paragraph()

heading("7.6. Rozmowy testowe (zapis sesji)", 2)
para(
    "Poniżej zamieszczono zapis sesji testowych przeprowadzonych podczas "
    "weryfikacji poprawności działania systemu. Screeny prezentują rzeczywiste "
    "działanie aplikacji w środowisku produkcyjnym (http://20.56.129.98/)."
)
doc.add_paragraph()
for i in range(1, 6):
    image_placeholder(f"Sesja testowa {i} — [opis testu]")
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 8. PODSUMOWANIE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("8. Podsumowanie", 1)

para(
    "Projekt Komunikator SR zrealizował wszystkie założone cele. "
    "Aplikacja działa w środowisku produkcyjnym na platformie Azure "
    "pod adresem http://20.56.129.98/ i umożliwia komunikację w czasie "
    "rzeczywistym przez WebSocket."
)
doc.add_paragraph()
para("Zrealizowane iteracje:", bold=True)
add_table(
    ["Iteracja", "Zakres", "Status"],
    [
        ["1", "Kick-off, wymagania, prezentacja wstępna", "✅"],
        ["2", "Szkielet projektu, Docker Compose, modele DB", "✅"],
        ["3", "Autentykacja JWT + bcrypt, formularze React", "✅"],
        ["4", "WebSocket, czat 1:1, typing, ACK, historia", "✅"],
        ["5", "Pokoje grupowe, broadcast, Socket.IO rooms", "✅"],
        ["6", "Status online, badge, toast, reconnect Infinity", "✅"],
        ["6.5", "Rola admina, panel admina, seed-admin CLI", "✅"],
        ["7.5", "Deployment Azure VM, nginx prod build, .env.prod", "✅"],
        ["8", "Dokumentacja techniczna", "✅"],
    ],
)
doc.add_paragraph()
para("Wymagania niefunkcjonalne:", bold=True)
add_table(
    ["Wymaganie", "Cel", "Osiągnięty"],
    [
        ["Opóźnienie RT", "< 100 ms (sieć lokalna)", "✅ WebSocket < 10 ms lokalnie"],
        ["Wielosesyjność", "Wiele kart/urządzeń", "✅ SessionStore: user_id → set(socket_ids)"],
        ["Bezpieczeństwo", "JWT + bcrypt, brak danych w URL", "✅"],
        ["Przenośność", "docker compose up", "✅"],
        ["Produkcja", "Dostęp przez internet", "✅ Azure VM, http://20.56.129.98/"],
    ],
)


# ── Zapis ─────────────────────────────────────────────────────────────────────

out = os.path.join(os.path.dirname(__file__), "Komunikator_SR_Dokumentacja.docx")
doc.save(out)
print(f"Zapisano: {out}")
